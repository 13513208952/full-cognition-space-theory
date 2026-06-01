#!/usr/bin/env python3
"""
将 full_paper.md 转换为学术格式 Word 文档
格式标准：中国学术期刊规范（GB/T 7714-2015，CSSCI期刊版式）

层级映射（pandoc markdown→docx）：
  #  (main title)    → Heading 1  → 黑体 16pt 居中
  ## (章节 1/2/…)    → Heading 2  → 黑体 14pt (一级标题)
  ###(小节 1.1/1.2…) → Heading 3  → 黑体 12pt (二级标题)

步骤：
1. build_reference_template() — python-docx 创建带学术样式的 reference.docx
2. run_pandoc()               — pandoc 将 Markdown+LaTeX 转 docx，挂载模板
"""

import os
import subprocess
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
INPUT_MD    = os.path.join(SCRIPT_DIR, "full_paper.md")
REF_DOCX    = os.path.join(SCRIPT_DIR, "reference.docx")
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, "full_paper.docx")


def apply_cn_font(style, cn_name, en_name, size_pt, bold=False, color=None, align=None):
    """为段落样式同时设置中英文字体、字号、加粗、颜色、对齐"""
    font = style.font
    font.name = en_name
    font.size = Pt(size_pt)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_name)
    rFonts.set(qn('w:ascii'), en_name)
    rFonts.set(qn('w:hAnsi'), en_name)
    if align is not None:
        style.paragraph_format.alignment = align


def set_para_spacing(style, before_pt=0, after_pt=0, line_spacing=None, first_line_pt=None):
    pf = style.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after  = Pt(after_pt)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)
    if first_line_pt is not None:
        pf.first_line_indent = Pt(first_line_pt)


def set_east_asian_font(run, cn_name):
    """为 run 补充东亚字体设置（用于页眉等）"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_name)


def build_reference_template():
    doc = Document()

    # ── 页面设置（A4，标准学术边距）────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    styles = doc.styles

    # ── Normal（正文，小四宋体，固定行距20磅，首行缩进2字符）──
    normal = styles['Normal']
    apply_cn_font(normal, '宋体', 'Times New Roman', 12)
    set_para_spacing(normal, before_pt=0, after_pt=0,
                     line_spacing=20, first_line_pt=24)

    # ── Heading 1 → 论文主标题（三号黑体16pt，居中）─────────
    # pandoc 将 markdown # 映射到 Heading 1
    h1 = styles['Heading 1']
    apply_cn_font(h1, '黑体', 'Times New Roman', 16,
                  bold=True, color=(0, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_para_spacing(h1, before_pt=12, after_pt=12)

    # ── Heading 2 → 一级节标题（四号黑体14pt）───────────────
    # pandoc 将 markdown ## 映射到 Heading 2
    h2 = styles['Heading 2']
    apply_cn_font(h2, '黑体', 'Times New Roman', 14,
                  bold=True, color=(0, 0, 0))
    set_para_spacing(h2, before_pt=18, after_pt=6)

    # ── Heading 3 → 二级小节标题（小四黑体12pt）─────────────
    # pandoc 将 markdown ### 映射到 Heading 3
    h3 = styles['Heading 3']
    apply_cn_font(h3, '黑体', 'Times New Roman', 12,
                  bold=True, color=(0, 0, 0))
    set_para_spacing(h3, before_pt=12, after_pt=6)

    # ── Heading 4 → 三级标题（五号宋体加粗10.5pt）───────────
    h4 = styles['Heading 4']
    apply_cn_font(h4, '宋体', 'Times New Roman', 10.5,
                  bold=True, color=(0, 0, 0))
    set_para_spacing(h4, before_pt=10, after_pt=4)

    # ── Block Text（定理/定义/证明块，五号宋体缩进）──────────
    try:
        block = styles['Block Text']
    except KeyError:
        block = styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
    apply_cn_font(block, '宋体', 'Times New Roman', 10.5)
    block.paragraph_format.left_indent  = Cm(0.8)
    block.paragraph_format.right_indent = Cm(0.8)
    set_para_spacing(block, before_pt=4, after_pt=4, line_spacing=16)

    # ── Abstract（五号宋体，左右缩进1.2cm）──────────────────
    try:
        abstract = styles['Abstract']
    except KeyError:
        abstract = styles.add_style('Abstract', WD_STYLE_TYPE.PARAGRAPH)
    apply_cn_font(abstract, '宋体', 'Times New Roman', 10.5)
    abstract.paragraph_format.left_indent  = Cm(1.2)
    abstract.paragraph_format.right_indent = Cm(1.2)
    set_para_spacing(abstract, before_pt=0, after_pt=6, line_spacing=16)

    # ── Caption（图注/表注，五号宋体居中）───────────────────
    try:
        caption = styles['Caption']
    except KeyError:
        caption = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    apply_cn_font(caption, '宋体', 'Times New Roman', 10.5,
                  color=(80, 80, 80), align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Title（YAML front matter title，备用，小二黑体居中）──
    title_style = styles['Title']
    apply_cn_font(title_style, '黑体', 'Times New Roman', 18,
                  bold=True, color=(0, 0, 0), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_para_spacing(title_style, before_pt=12, after_pt=12)

    # ── 页眉（小五宋体居中）─────────────────────────────────
    header = sec.header
    hp = header.paragraphs[0]
    hp.text = "全认识空间理论：AI认知极限的数学证明"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = hp.runs[0]
    run_h.font.name = 'Times New Roman'
    run_h.font.size = Pt(9)
    set_east_asian_font(run_h, '宋体')

    # ── 页脚（居中页码，小五）───────────────────────────────
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = fp.add_run()
    run_f.font.size = Pt(9)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_f._element.append(fldChar1)
    run_f._element.append(instrText)
    run_f._element.append(fldChar2)

    doc.save(REF_DOCX)
    print(f"  参考模板已生成: {REF_DOCX}")


def run_pandoc():
    cmd = [
        "pandoc",
        INPUT_MD,
        "--from", "markdown+tex_math_dollars+tex_math_single_backslash",
        "--to", "docx",
        "--reference-doc", REF_DOCX,
        "--wrap=none",
        "-o", OUTPUT_DOCX,
    ]
    print(f"  运行 pandoc...")
    result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
    if result.returncode != 0:
        print(f"  [错误] pandoc 退出码 {result.returncode}")
        print(result.stderr)
    else:
        print(f"  转换成功: {OUTPUT_DOCX}")
        if result.stderr:
            print(f"  警告: {result.stderr[:300]}")


if __name__ == "__main__":
    print("第一步：生成学术样式参考模板（中国学术期刊规范）...")
    build_reference_template()
    print("第二步：pandoc 转换 Markdown → Word（不重复编号）...")
    run_pandoc()
    size = os.path.getsize(OUTPUT_DOCX) // 1024
    print(f"\n完成。输出文件大小：{size} KB")
    print(f"文件路径：{OUTPUT_DOCX}")
